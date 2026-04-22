# author: jf
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

from app.domain.exceptions.rag_exceptions import FileParseError, UnsupportedFileTypeError
from app.domain.models.rag_document import ExtractedDocument


class FileParserAdapter:
    _TEXT_EXTENSIONS = {".txt", ".md"}
    _PDF_EXTENSIONS = {".pdf"}
    _DOCX_EXTENSIONS = {".docx"}
    _WORDPROCESSING_NAMESPACE = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    def parse(
        self,
        file_bytes: bytes,
        file_name: str,
        content_type: str | None = None,
    ) -> ExtractedDocument:
        safe_name = (file_name or "").strip() or "document"
        extension = Path(safe_name).suffix.lower()
        safe_content_type = (content_type or "").strip() or "application/octet-stream"
        _log_parser(
            "开始解析文档文件",
            file_name=safe_name,
            extension=extension or "none",
            content_type=safe_content_type,
            size_bytes=len(file_bytes),
        )

        if not file_bytes:
            raise FileParseError("上传文件不能为空")

        if extension in self._TEXT_EXTENSIONS:
            content = self._decode_text(file_bytes)
        elif extension in self._PDF_EXTENSIONS:
            content = self._parse_pdf(file_bytes)
        elif extension in self._DOCX_EXTENSIONS:
            content = self._parse_docx(file_bytes)
        else:
            raise UnsupportedFileTypeError(f"暂不支持的文件类型: {extension or safe_content_type}")

        normalized = content.strip()
        if not normalized:
            raise FileParseError(f"文件 {safe_name} 未解析到可入库内容")

        source_id = Path(safe_name).stem.strip() or safe_name
        _log_parser("文档解析完成", file_name=safe_name, parsed_chars=len(normalized), source_id=source_id)
        return ExtractedDocument(
            source_id=source_id,
            original_filename=safe_name,
            original_content_type=safe_content_type,
            source_type="document",
            ingest_source="text_document",
            content=normalized,
        )

    @staticmethod
    def _decode_text(file_bytes: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gbk"):
            try:
                _log_parser("尝试按编码解析文本", encoding=encoding)
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        _log_parser("文本编码兜底解析", encoding="utf-8-replace")
        return file_bytes.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_pdf(file_bytes: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise FileParseError("缺少 pypdf 依赖，无法解析 PDF 文件") from exc

        try:
            reader = PdfReader(BytesIO(file_bytes))
            texts = [(page.extract_text() or "").strip() for page in reader.pages]
            _log_parser("PDF 解析完成", page_count=len(reader.pages))
        except Exception as exc:
            raise FileParseError(f"PDF 解析失败: {exc.__class__.__name__}: {exc}") from exc
        return "\n\n".join(text for text in texts if text)

    @classmethod
    def _parse_docx(cls, file_bytes: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise FileParseError("缺少 python-docx 依赖，无法解析 DOCX 文件") from exc

        try:
            document = Document(BytesIO(file_bytes))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            if paragraphs:
                _log_parser("DOCX 通过 python-docx 解析成功", paragraph_count=len(paragraphs))
                return "\n\n".join(paragraphs)
        except Exception as exc:
            _log_parser("DOCX 主解析失败，尝试 OOXML 兜底", error=str(exc))
            fallback_text = cls._parse_docx_xml_fallback(file_bytes)
            if fallback_text:
                return fallback_text
            raise FileParseError(f"DOCX 解析失败: {exc.__class__.__name__}: {exc}") from exc

        fallback_text = cls._parse_docx_xml_fallback(file_bytes)
        if fallback_text:
            return fallback_text
        raise FileParseError("DOCX 未解析到可用文本内容")

    @classmethod
    def _parse_docx_xml_fallback(cls, file_bytes: bytes) -> str:
        try:
            with ZipFile(BytesIO(file_bytes)) as archive:
                document_xml = archive.read("word/document.xml")
        except (BadZipFile, KeyError, ValueError):
            return ""

        try:
            root = ElementTree.fromstring(document_xml)
        except ElementTree.ParseError:
            return ""

        paragraphs: list[str] = []
        for paragraph in root.findall(".//w:body//w:p", cls._WORDPROCESSING_NAMESPACE):
            fragments: list[str] = []
            for node in paragraph.iter():
                tag_name = node.tag.rsplit("}", 1)[-1]
                if tag_name == "t" and node.text:
                    fragments.append(node.text)
                elif tag_name == "tab":
                    fragments.append("\t")
                elif tag_name in {"br", "cr"}:
                    fragments.append("\n")
            paragraph_text = "".join(fragments).strip()
            if paragraph_text:
                paragraphs.append(paragraph_text)

        if paragraphs:
            _log_parser("DOCX OOXML 段落兜底解析成功", paragraph_count=len(paragraphs))
            return "\n\n".join(paragraphs)

        text_nodes = [
            (node.text or "").strip()
            for node in root.findall(".//w:t", cls._WORDPROCESSING_NAMESPACE)
            if (node.text or "").strip()
        ]
        if text_nodes:
            _log_parser("DOCX OOXML 文本节点兜底解析成功", node_count=len(text_nodes))
        return "\n".join(text_nodes)


def _log_parser(message: str, **extra: object) -> None:
    parts = [f"[知识库上传][文件解析] {message}"]
    for key, value in extra.items():
        parts.append(f"{key}={value}")
    print(" ".join(parts), flush=True)
